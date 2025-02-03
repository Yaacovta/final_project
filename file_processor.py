# import PyPDF2
from docx import Document
from fpdf import FPDF
import sys
import re
import pdfplumber
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# os.system("chcp 65001")  # Change encoding to UTF-8 (Hebrew)
sys.stdout.reconfigure(encoding='utf-8')  # Change encoding to UTF-8 (Hebrew)


def reverse_hebrew_word(word):
    """
    Reverses the order of letters in a word if it is in Hebrew.
    """
    if re.search(r'[\u0590-\u05FF]', word):  # Check if the word contains Hebrew characters
        return word[::-1]  # Reverse the order of letters
    return word


def reverse_words_in_line(line):
    """
    Reverses the letters in Hebrew words and flips the order of words in the line.
    """
    words = line.split()  # Split the line into words
    processed_words = [reverse_hebrew_word(word) for word in words]  # Process each word
    reversed_words = list(reversed(processed_words))  # Reverse the order of words in the line
    return ' '.join(reversed_words)  # Join the words back into a line


def process_text(text):
    """
    Processes the entire text:
    - Reverses the order of words in each line.
    - Reverses the letters in Hebrew words.
    """
    lines = text.split('\n')  # Split the text into lines
    processed_lines = [reverse_words_in_line(line) for line in lines]  # Process each line
    return '\n'.join(processed_lines)  # Join the processed lines into a complete text


def extract_text_from_pdf(file_path):
    """
    Extracts text from a PDF file and processes it.
    """
    try:
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + '\n'  # Append text from each page with a newline
        return process_text(text)
    except Exception as e:
        print(f"An error occurred while reading the text from the PDF: {e}")
        return None


def extract_text_from_docx(file_path):
    """
    Extracts text from a DOCX file and processes it.
    """
    try:
        doc = Document(file_path)  # Open the DOCX file
        text = [paragraph.text for paragraph in doc.paragraphs]  # Extract text from each paragraph
        return "\n".join(text)  # Combine paragraphs into a single string
    except Exception as e:
        print(f"An error occurred while reading the text from the DOCX: {e}")
        return None


def extract_text(file_path):
    """
    Determines the file type and extracts text accordingly.

    If the file is a PDF, it uses the extract_text_from_pdf function.
    If the file is a DOCX, it uses the extract_text_from_docx function.
    """
    if file_path.endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    if file_path.endswith(".docx"):
        return extract_text_from_docx(file_path)


ex = extract_text("examm.pdf")
# print(ex)

ex2 = extract_text("examm2.docx")
# print(ex2)


txt_file = "exist_exercise.txt"
with open(txt_file, "w", encoding="utf-8") as file:
    file.write(ex2)


def reverse_line_words_and_letters(line):
    """
    Reverse the order of words in a line and reverse letters in Hebrew words only.
    :param line: A single line of text.
    :return: The line with reversed word order and reversed letters in Hebrew words.
    """
    words = line.split(" ")  # Split the line into words
    processed_words = []
    for word in reversed(words):  # Reverse the word order
        # Reverse letters in Hebrew words
        if any("\u0590" <= char <= "\u05FF" for char in word):
            processed_words.append(word[::-1])  # Reverse the Hebrew word
        else:
            processed_words.append(word)  # Keep English words as they are
    return " ".join(processed_words)  # Combine the processed words back into a line


def save_as_pdf(output_text, name_of_new_exercise):
    """
    Save the provided text content as a PDF file, reversing the word order and reversing Hebrew letters.
    :param output_text: The text to save into the PDF.
    :param name_of_new_exercise: The name of the PDF file to create (without extension).
    """
    pdf = FPDF()
    pdf.add_page()

    # Load a font that supports Hebrew and English
    pdf.add_font('Arial', '', 'Arial_font/arial.ttf', uni=True)
    pdf.set_font("Arial", size=12)

    # Split the text into lines
    lines = output_text.split("\n")
    for line in lines:
        # Reverse the word order and reverse Hebrew letters
        processed_line = reverse_line_words_and_letters(line)
        # Use multi_cell to wrap text and align to the right
        pdf.multi_cell(0, 10, processed_line, align="R")

    # Save the PDF
    pdf_file = f"{name_of_new_exercise}.pdf"
    pdf.output(pdf_file)
    print(f"PDF saved as: {pdf_file}")


def save_as_docx(output_text, name_of_new_exercise):
    """
    Save the text as a DOCX file, with proper RTL handling for Hebrew text.
    """
    doc = Document()

    doc.add_paragraph(output_text)

    doc.save(f"{name_of_new_exercise}.docx")
    print(f"Saved as DOCX: {name_of_new_exercise}.docx")


def file_type_of_the_new_exercise(output_text, name_of_new_exercise, output_format):
    """
    Save the text as either PDF or DOCX based on the output format.
    """
    if output_format == "P":
        save_as_pdf(output_text, name_of_new_exercise)
    elif output_format == "D":
        save_as_docx(output_text, name_of_new_exercise)
    else:
        raise ValueError("Unsupported format. Please choose 'P' for PDF, 'D' for DOCX.")
